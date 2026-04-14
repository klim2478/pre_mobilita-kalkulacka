import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

# Nastavení stránky
st.set_page_config(page_title="PRE Nabídkový generátor", layout="wide")

# --- DATA A LOGIKA MODELŮ ---
models = {
    "Model 3.1 - Realizace na stávající přívod (21 míst)": {
        "popis": "Realizace na stávající přívod, kabelové vedení: 21 míst, 50 % penetrace (bez úpravy HDV).",
        "rozsiritelnost": "Nízká – omezeno kapacitou stávajícího jističe a místem v kabelových trasách.",
        "pater_svj": 450000,
        "hdv_navyseni": 0,
        "zakaznik_celkem": 1197000,
        "wallbox_ks": 23800,
        "mist": 21,
        "schema": "schema_kabel.png"
    },
    "Model 3.2 - Přípojnicové řešení (100 míst)": {
        "popis": "Realizace 100 míst, 100 % penetrace – přípojnicové vedení (obsahuje nové HDV a navýšení kapacity).",
        "rozsiritelnost": "Vysoká – páteřní systém umožňuje snadné připojení jakéhokoliv stání v budoucnu bez nutnosti dalších stavebních úprav.",
        "pater_svj": 1250000,
        "hdv_navyseni": 580000,
        "zakaznik_celkem": 3500000,
        "wallbox_ks": 23800,
        "mist": 100,
        "schema": "schema_pripajnice.png"
    },
    "Model 3.3 - Malá realizace (10 míst)": {
        "popis": "Realizace 10 míst, 100 % penetrace, kabelové vedení (obsahuje nové HDV a navýšení kapacity).",
        "rozsiritelnost": "Střední – díky novému HDV je kapacita dostatečná, ale kabelové trasy mohou být při dalším rozšiřování komplikované.",
        "pater_svj": 290000,
        "hdv_navyseni": 500000,
        "zakaznik_celkem": 330000,
        "wallbox_ks": 23800,
        "mist": 10,
        "schema": "schema_kabel.png"
    }
}

# --- SIDEBAR VSTUPY ---
st.sidebar.header("📝 Údaje pro nabídku")
projekt_nazev = st.sidebar.text_input("Název projektu / SVJ:", "SVJ ")
sel_name = st.sidebar.selectbox("Varianta realizace:", list(models.keys()))
podil_svj = st.sidebar.slider("Podíl SVJ na zákaznických rozvodech v garážích (%)", 0, 100, 0)
dph_check = st.sidebar.checkbox("Zobrazit ceny s DPH (12 %)", value=False)

# Výchozí parametry pro zvolený model
m = models[sel_name]
coeff = 1.12 if dph_check else 1.0
tax_label = "Včetně DPH 12 %" if dph_check else "Bez DPH"

# --- VÝPOČTY ---
# Náklady SVJ (Páteř + HDV)
naklady_svj_pater = (m["pater_svj"] + m["hdv_navyseni"]) * coeff
# Příspěvek SVJ na zákaznické rozvody (dle slideru)
extra_z_garazi = (m["zakaznik_celkem"] * (podil_svj / 100)) * coeff
# Celková investice SVJ
celkem_svj_investice = naklady_svj_pater + extra_z_garazi

# Náklad na jednoho uživatele (rozvody po odečtení podílu SVJ + Wallbox)
naklad_uzivatel_rozvody = ((m["zakaznik_celkem"] * (1 - podil_svj / 100)) / m["mist"]) * coeff
celkem_uzivatel = naklad_uzivatel_rozvody + (m["wallbox_ks"] * coeff)

# --- ZOBRAZENÍ V APLIKACI ---
st.title(f"Indikativní kalkulace: {projekt_nazev}")
st.subheader(f"Cenová hladina: {tax_label}")

col1, col2 = st.columns(2)
with col1:
    st.metric("Celková investice SVJ", f"{celkem_svj_investice:,.0f} Kč".replace(",", " "))
    st.info(f"**Rozšiřitelnost:** {m['rozsiritelnost']}")
with col2:
    st.metric("Náklad na 1 uživatele (vč. WB)", f"{celkem_uzivatel:,.0f} Kč".replace(",", " "))
    st.warning("Doporučeno: Modul COMFORT (bez starostí pro SVJ)")

st.divider()

# Příprava dat pro tabulku
table_data = [
    {
        "Položka": "Páteřní rozvody a nové HDV",
        "Hradí SVJ": f"{naklady_svj_pater:,.0f} Kč".replace(",", " "),
        "Hradí 1 Uživatel": "0 Kč"
    },
    {
        "Položka": f"Zákaznické rozvody (příspěvek SVJ {podil_svj} %)",
        "Hradí SVJ": f"{extra_z_garazi:,.0f} Kč".replace(",", " "),
        "Hradí 1 Uživatel": f"{naklad_uzivatel_rozvody:,.0f} Kč".replace(",", " ")
    },
    {
        "Položka": "Wallbox vč. instalace",
        "Hradí SVJ": "0 Kč",
        "Hradí 1 Uživatel": f"{(m['wallbox_ks'] * coeff):,.0f} Kč".replace(",", " ")
    },
    {
        "Položka": "CELKEM (součet položek)",
        "Hradí SVJ": f"{celkem_svj_investice:,.0f} Kč".replace(",", " "),
        "Hradí 1 Uživatel": f"{celkem_uzivatel:,.0f} Kč".replace(",", " ")
    }
]

# Zobrazení interaktivní tabulky
st.table(pd.DataFrame(table_data))

# --- GENERÁTOR WORD DOKUMENTU ---
def create_word():
    doc = Document()
    
    # Styl písma pro celý dokument
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Titulní strana
    title = doc.add_heading('Indikativní nákladovost emobility v SVJ/BD dle vzorové instalace', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Projekt: {projekt_nazev}\nVarianta: {sel_name}\nCenová hladina: {tax_label}')
    run.bold = True

    # 2. Argumentace a výhody
    doc.add_heading('Klíčové výhody řešení v Modulu COMFORT', level=1)
    
    para1 = doc.add_paragraph()
    para1.add_run('Bezstarostné rozúčtování: ').bold = True
    para1.add_run('V modulu COMFORT fakturuje PRE spotřebu přímo koncovému uživateli na základě vlastního měření. SVJ odpadá administrativa spojená s rozúčtováním nákladů mezi nájemníky.')

    para2 = doc.add_paragraph()
    para2.add_run('Ochrana proti přetížení (DLM): ').bold = True
    para2.add_run('Systém dynamického řízení výkonu monitoruje aktuální odběr celého domu a reguluje nabíjení tak, aby nedošlo k výpadku hlavního jističe ani při špičce v bytech.')

    # 3. Technický popis
    doc.add_heading('Technická koncepce', level=1)
    doc.add_paragraph(f'Popis zvoleného řešení: {m["popis"]}')
    doc.add_paragraph(f'Rozšiřitelnost: {m["rozsiritelnost"]}')

    # 4. Ekonomická tabulka
    doc.add_heading('Ekonomické rozdělení nákladů', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Hlavička tabulky
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Položka'
    hdr_cells[1].text = 'Hradí SVJ (Fond oprav)'
    hdr_cells[2].text = 'Hradí 1 Uživatel'

    # Plnění dat do tabulky
    for i, row_item in enumerate(table_data):
        row_cells = table.add_row().cells
        
        # Pokud jde o poslední řádek (Součet), nastavíme tučné písmo
        if i == len(table_data) - 1:
            for cell in row_cells:
                # Přístup k odstavci v buňce pro nastavení formátu
                p = cell.paragraphs[0]
                run = p.add_run()
                run.bold = True
                
        row_cells[0].text = row_item["Položka"]
        row_cells[1].text = row_item["Hradí SVJ"]
        row_cells[2].text = row_item["Hradí 1 Uživatel"]

    # 5. Schéma zapojení (pokud soubor existuje)
    if os.path.exists(m["schema"]):
        doc.add_heading('Vzorové schéma zapojení', level=1)
        doc.add_picture(m["schema"], width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Uložení do paměti a vrácení bytů
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

# --- TLAČÍTKO KE STAŽENÍ ---
st.divider()
if st.button("📄 Vygenerovat prezentaci pro SVJ (.docx)"):
    with st.spinner('Generuji dokument...'):
        doc_bytes = create_word()
        st.download_button(
            label="📥 Stáhnout hotovou prezentaci",
            data=doc_bytes,
            file_name=f"Nabidka_PRE_{projekt_nazev.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
